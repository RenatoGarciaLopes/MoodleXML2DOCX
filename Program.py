import os
import xml.etree.ElementTree as ET
from docx import Document
from html import unescape
import re

input_folder = "BQS"
output_folder = "BQS-CONV"
os.makedirs(output_folder, exist_ok=True)

class HTMLCleaner:
    @staticmethod
    def clean(html):
        html = html.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
        html = re.sub(r'<\/?p>', '\n', html)
        html = re.sub(r'&nbsp;', ' ', html)
        html = re.sub(r'<[^>]+>', '', html)
        html = unescape(html)
        html = re.sub(r'\n\s+', '\n', html)
        html = re.sub(r'\n{2,}', '\n\n', html)
        return html.strip()

def process_question_to_docx(input_file, output_file):
    tree = ET.parse(input_file)
    root = tree.getroot()
    
    doc = Document()
    doc.add_heading("Banco de Questões", level=1)

    for question in root.findall('question'):
        if question.get('type') != 'multichoice':
            continue

        # Extrair texto da questão
        question_text_elem = question.find("./questiontext/text")
        question_text = (
            HTMLCleaner.clean(question_text_elem.text)
            if question_text_elem is not None and question_text_elem.text else ""
        )

        if not question_text:
            continue

        # Adicionar questão
        doc.add_paragraph("Questão:", style='Heading 2')
        doc.add_paragraph(question_text)

        # Adiciona as alternativas
        alternatives = question.findall("./answer")
        if alternatives:
            for i, answer in enumerate(alternatives):
                answer_text_elem = answer.find("text")
                answer_text = (
                    HTMLCleaner.clean(answer_text_elem.text)
                    if answer_text_elem is not None and answer_text_elem.text else ""
                )
                
                if answer_text:  # Adiciona apenas alternativas válidas
                    letter = chr(97 + i)  # Gera 'a', 'b', 'c', etc.
                    # Verifica se é a resposta correta
                    fraction = answer.get("fraction", "0")
                    is_correct = (fraction == "100")
                    
                    # Cria um parágrafo com estilo de lista
                    p = doc.add_paragraph(style="List Bullet")
                    # Adiciona a letra da alternativa (sem formatação)
                    run_letter = p.add_run(f"{letter}) ")
                    # Adiciona o texto da resposta (negrito se for correta)
                    run_text = p.add_run(answer_text)
                    if is_correct:
                        run_letter.bold = True  # Letra em negrito
                        run_text.bold = True    # Texto em negrito

    doc.save(output_file)

# Processar arquivos
try:
    for filename in os.listdir(input_folder):
        if filename.endswith(".xml"):
            input_path = os.path.join(input_folder, filename)
            output_path = os.path.join(output_folder, filename.replace(".xml", ".docx"))
            process_question_to_docx(input_path, output_path)
    print(f"Conversão concluída! Arquivos em: '{output_folder}'")
except Exception as e:
    print(f"Erro: {str(e)}")

input("Pressione Enter para sair...")